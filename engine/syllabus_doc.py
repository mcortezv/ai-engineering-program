# -*- coding: utf-8 -*-
"""
Genera "Programa de AI - Hyperlabs.docx" y su exportacion a PDF.

Sistema de diseno tomado de hyperlabs-press/paper/print.css (kit de prensa de
HyperLabs), con una desviacion deliberada: los titulos de seccion NO llevan la
regla morada inferior que usa el CSS original (h2 border-bottom). El morado se
conserva como acento en todo lo demas.

El contenido vive en syllabus_p1..p5.py y syllabus_anexos.py.

    python build_syllabus.py          -> genera docx y pdf
    python build_syllabus.py --docx   -> solo docx
"""

import os
import re
import subprocess
import sys

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
DIST = os.path.join(os.path.dirname(HERE), "dist")
DOCX_PATH = os.path.join(DIST, "Programa de AI - Hyperlabs.docx")
PDF_PATH = os.path.join(DIST, "Programa de AI - Hyperlabs.pdf")


# ============================================================ design tokens ==
# Equivalentes de las variables :root de print.css

INK = RGBColor(0x16, 0x16, 0x1A)        # --ink
INK2 = RGBColor(0x3D, 0x3D, 0x46)       # --ink2
INK3 = RGBColor(0x6D, 0x6D, 0x78)       # --ink3
ACC = RGBColor(0x5B, 0x3B, 0xD6)        # --acc
OK = RGBColor(0x0F, 0x7A, 0x52)         # --ok
WARN = RGBColor(0x8A, 0x5A, 0x06)       # --warn
DANGER = RGBColor(0xB9, 0x1C, 0x1C)     # --danger
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HAIR = "D8D8E0"          # --hair
ACC_HEX = "5B3BD6"
ACC_SOFT = "F1EDFF"      # --acc-soft
CODE_BG = "F7F7FA"       # pre
INLINE_BG = "F4F4F7"     # code
TH_BG = "FAFAFC"         # th
ZEBRA = "FBFBFD"         # tbody tr:nth-child(even)
INK3_HEX = "6D6D78"

# Las tintas suaves de ok/warn/danger no estan en el CSS (solo define
# --acc-soft); se derivan con la misma logica: el acento al ~8% sobre blanco.
BOX_STYLES = {
    "nota": (ACC_SOFT, ACC_HEX, ACC),
    "correccion": ("E9F5F0", "0F7A52", OK),
    "volatil": ("FBF3E3", "8A5A06", WARN),
    "alerta": ("FDECEC", "B91C1C", DANGER),
}

# --serif / --sans / --mono, resueltos a los fallbacks instalados en el equipo.
SERIF = "Georgia"
SANS = "Segoe UI"
MONO = "Consolas"

BODY_PT = 9.9

# Grosores de borde en octavos de punto (unidad de w:sz).
SZ_HAIR = 4      # ~0.5 pt
SZ_TH = 8        # ~1.0 pt
SZ_ACCENT = 18   # ~2.3 pt


# ------------------------------------------------------------- primitivas ----

def _shade(element, hex_fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    element.append(shd)


def _char_spacing(run, points):
    """letter-spacing. w:spacing va en veinteavos de punto."""
    rPr = run._r.get_or_add_rPr()
    node = OxmlElement("w:spacing")
    node.set(qn("w:val"), str(int(round(points * 20))))
    rPr.append(node)


def _para_border(pPr, edge, hex_color, size, space=8):
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    node = OxmlElement("w:" + edge)
    node.set(qn("w:val"), "single")
    node.set(qn("w:sz"), str(size))
    node.set(qn("w:space"), str(space))
    node.set(qn("w:color"), hex_color)
    pBdr.append(node)


def _cell_margins(cell, top=90, start=130, bottom=90, end=130):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement("w:" + tag)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


_EDGE_ORDER = ("top", "start", "left", "bottom", "end", "right", "insideH", "insideV")


def _borders(element, spec, tag):
    """spec: {edge: (hex, size) | None}. tag es 'w:tblBorders' o 'w:tcBorders'."""
    node = OxmlElement(tag)
    for edge in _EDGE_ORDER:
        if edge not in spec:
            continue
        child = OxmlElement("w:" + edge)
        value = spec[edge]
        if value is None:
            child.set(qn("w:val"), "none")
            child.set(qn("w:sz"), "0")
        else:
            hex_color, size = value
            child.set(qn("w:val"), "single")
            child.set(qn("w:sz"), str(size))
            child.set(qn("w:space"), "0")
            child.set(qn("w:color"), hex_color)
        node.append(child)
    element.append(node)


def _table_borders(table, spec):
    _borders(table._tbl.tblPr, spec, "w:tblBorders")


def _cell_borders(cell, spec):
    _borders(cell._tc.get_or_add_tcPr(), spec, "w:tcBorders")


def _no_split(table):
    """Evita que una fila se parta entre paginas (page-break-inside:avoid)."""
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        node = OxmlElement("w:cantSplit")
        trPr.append(node)


def _field(paragraph, instruction):
    for kind in ("begin",):
        run = paragraph.add_run()._r
        node = OxmlElement("w:fldChar")
        node.set(qn("w:fldCharType"), kind)
        run.append(node)

    run = paragraph.add_run()._r
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run.append(instr)

    run = paragraph.add_run()._r
    node = OxmlElement("w:fldChar")
    node.set(qn("w:fldCharType"), "separate")
    run.append(node)

    placeholder = paragraph.add_run("")

    run = paragraph.add_run()._r
    node = OxmlElement("w:fldChar")
    node.set(qn("w:fldCharType"), "end")
    run.append(node)
    return placeholder


# -------------------------------------------------------- texto enriquecido --

_INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*\n]+?\*)")


def add_rich(paragraph, text, size=None, color=None, bold=False, italic=False,
             font=None):
    base = size or BODY_PT
    parts = []
    for i, line in enumerate(text.split("\n")):
        if i:
            parts.append("\n")
        parts.extend(_INLINE.split(line))

    for part in parts:
        if not part:
            continue
        if part == "\n":
            paragraph.add_run().add_break()
            continue

        is_code = False
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            # code inline: fondo gris muy claro y texto en el acento.
            run = paragraph.add_run(part[1:-1])
            run.font.name = MONO
            run.font.size = Pt(base - 1.4)
            run.font.color.rgb = ACC
            _shade(run._r.get_or_add_rPr(), INLINE_BG)
            is_code = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)

        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if not is_code:
            if font:
                run.font.name = font
            if size:
                run.font.size = Pt(size)
            if color is not None:
                run.font.color.rgb = color
    return paragraph


# ----------------------------------------------------------------- estilos ---

def build_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = SERIF
    normal.font.size = Pt(BODY_PT)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.52
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Idioma es-MX: activa la particion silabica y quita los subrayados rojos.
    rPr = normal.element.get_or_add_rPr()
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "es-MX")
    rPr.append(lang)

    def new(name, base="Normal"):
        st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = styles[base]
        st.quick_style = False
        return st

    # --- Heading 1: titulo de parte (pagina divisoria) y titulo de anexo.
    h1 = styles["Heading 1"]
    h1.font.name = SANS
    h1.font.size = Pt(21)
    h1.font.bold = True
    h1.font.color.rgb = INK
    h1.paragraph_format.space_before = Pt(2)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --- Heading 2: modulo. Sin la regla morada inferior del CSS original.
    h2 = styles["Heading 2"]
    h2.font.name = SANS
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = INK
    h2.paragraph_format.space_before = Pt(20)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --- Heading 3: subsecciones numeradas (9.1, 13.1, B.1.1) y de anexo.
    h3 = styles["Heading 3"]
    h3.font.name = SANS
    h3.font.size = Pt(11.6)
    h3.font.bold = True
    h3.font.color.rgb = ACC
    h3.paragraph_format.space_before = Pt(17)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --- Etiqueta de seccion fija del modulo. Versaditas gris con filete
    #     superior: da ritmo sin recurrir a una regla de color.
    sec = new("SeccionModulo")
    sec.font.name = SANS
    sec.font.size = Pt(8)
    sec.font.bold = True
    sec.font.all_caps = True
    sec.font.color.rgb = INK3
    sec.paragraph_format.space_before = Pt(16)
    sec.paragraph_format.space_after = Pt(6)
    sec.paragraph_format.keep_with_next = True
    sec.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_border(sec.element.get_or_add_pPr(), "top", HAIR, SZ_HAIR, space=6)
    # Fuera del indice: si no, 80 lineas de ruido.
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), "9")
    sec.element.get_or_add_pPr().append(outline)

    # --- Subtitulos del enfoque didactico. Son el elemento de estructura mas
    #     frecuente, asi que llevan el morado del h3 del CSS de referencia.
    sub = new("SubEnfoque")
    sub.font.name = SANS
    sub.font.size = Pt(10.4)
    sub.font.bold = True
    sub.font.color.rgb = ACC
    sub.paragraph_format.space_before = Pt(13)
    sub.paragraph_format.space_after = Pt(3)
    sub.paragraph_format.keep_with_next = True
    sub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --- Etiqueta de campo dentro de fichas y recuadros.
    lab = new("Etiqueta")
    lab.font.name = SANS
    lab.font.size = Pt(7.4)
    lab.font.bold = True
    lab.font.all_caps = True
    lab.font.color.rgb = INK3
    lab.paragraph_format.space_before = Pt(0)
    lab.paragraph_format.space_after = Pt(2)
    lab.paragraph_format.keep_with_next = True
    lab.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    celda = new("Celda")
    celda.font.name = SANS
    celda.font.size = Pt(8.5)
    celda.paragraph_format.space_after = Pt(0)
    celda.paragraph_format.line_spacing = 1.3
    celda.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    cuerpo_caja = new("CuerpoCaja")
    cuerpo_caja.font.size = Pt(9.5)
    cuerpo_caja.paragraph_format.space_after = Pt(0)
    cuerpo_caja.paragraph_format.line_spacing = 1.45

    code = new("Codigo")
    code.font.name = MONO
    code.font.size = Pt(7.9)
    code.font.color.rgb = INK2
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.48
    code.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    cap = new("Pie")
    cap.font.name = SANS
    cap.font.size = Pt(8.2)
    cap.font.color.rgb = INK3
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(11)
    cap.paragraph_format.line_spacing = 1.42
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_border(cap.element.get_or_add_pPr(), "top", HAIR, 2, space=4)

    # Listas propias: los estilos "List Bullet"/"List Number" de Word arrastran
    # una sangria grande y una numeracion compartida que continua entre modulos.
    # CSS: ul,ol{padding-left:5.5mm}
    num = new("ListaNum")
    num.paragraph_format.left_indent = Cm(0.62)
    num.paragraph_format.first_line_indent = Cm(-0.62)
    num.paragraph_format.space_after = Pt(3)
    num.paragraph_format.line_spacing = 1.4
    num.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    num.paragraph_format.tab_stops.add_tab_stop(Cm(0.62))

    b1 = new("Vineta")
    b1.paragraph_format.left_indent = Cm(0.55)
    b1.paragraph_format.first_line_indent = Cm(-0.55)
    b1.paragraph_format.space_after = Pt(3)
    b1.paragraph_format.line_spacing = 1.4
    b1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    b1.paragraph_format.tab_stops.add_tab_stop(Cm(0.55))

    b2 = new("Vineta2")
    b2.font.size = Pt(9.4)
    b2.paragraph_format.left_indent = Cm(1.15)
    b2.paragraph_format.first_line_indent = Cm(-0.53)
    b2.paragraph_format.space_after = Pt(2)
    b2.paragraph_format.line_spacing = 1.36
    b2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    b2.paragraph_format.tab_stops.add_tab_stop(Cm(1.15))


def enable_hyphenation(doc):
    settings = doc.settings.element
    node = OxmlElement("w:autoHyphenation")
    node.set(qn("w:val"), "1")
    settings.append(node)
    zone = OxmlElement("w:hyphenationZone")
    zone.set(qn("w:val"), "284")   # 0.5 cm
    settings.append(zone)


# ------------------------------------------------------------ componentes ----

def text_width(doc):
    section = doc.sections[0]
    return section.page_width - section.left_margin - section.right_margin


def add_box(doc, kind, title, paragraphs):
    """Equivalente de #abstract-box: tinte suave y filete de color a la izquierda."""
    fill, border_hex, title_color = BOX_STYLES[kind]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _table_borders(table, {
        "left": (border_hex, SZ_ACCENT), "start": (border_hex, SZ_ACCENT),
        "top": None, "bottom": None, "right": None, "end": None,
        "insideH": None, "insideV": None,
    })
    # Sin cantSplit a proposito: una caja alta que no puede partirse arrastra su
    # encabezado a la pagina siguiente y deja medio folio en blanco.

    cell = table.cell(0, 0)
    cell.width = text_width(doc)
    _shade(cell._tc.get_or_add_tcPr(), fill)
    _cell_margins(cell, top=150, bottom=150, start=200, end=200)

    target = cell.paragraphs[0]
    if title:
        target.style = doc.styles["Etiqueta"]
        run = target.add_run(title.upper())
        run.font.color.rgb = title_color
        _char_spacing(run, 0.52)
        target = None

    for i, text in enumerate(paragraphs):
        p = target if (target is not None and i == 0) else cell.add_paragraph()
        p.style = doc.styles["CuerpoCaja"]
        p.paragraph_format.space_after = Pt(5) if i < len(paragraphs) - 1 else Pt(0)
        add_rich(p, text, size=9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_code_block(doc, text):
    """Equivalente de pre: fondo gris, filete morado a la izquierda."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _table_borders(table, {
        "left": (ACC_HEX, SZ_ACCENT), "start": (ACC_HEX, SZ_ACCENT),
        "top": (HAIR, SZ_HAIR), "bottom": (HAIR, SZ_HAIR),
        "right": (HAIR, SZ_HAIR), "end": (HAIR, SZ_HAIR),
        "insideH": None, "insideV": None,
    })
    _no_split(table)

    cell = table.cell(0, 0)
    cell.width = text_width(doc)
    _shade(cell._tc.get_or_add_tcPr(), CODE_BG)
    _cell_margins(cell, top=130, bottom=130, start=180, end=140)

    lines = text.split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.style = doc.styles["Codigo"]
        p.add_run(line if line.strip() else " ")

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_table(doc, headers, rows, widths=None, caption=None):
    """Tabla de paper: solo filetes horizontales, cabecera en versaditas."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _table_borders(table, {
        "top": None, "left": None, "start": None, "right": None, "end": None,
        "insideV": None,
        "bottom": (HAIR, SZ_HAIR), "insideH": (HAIR, SZ_HAIR),
    })

    header = table.rows[0]
    for i, text in enumerate(headers):
        cell = header.cells[i]
        _shade(cell._tc.get_or_add_tcPr(), TH_BG)
        _cell_borders(cell, {"bottom": (INK3_HEX, SZ_TH)})
        _cell_margins(cell, top=70, bottom=70)
        p = cell.paragraphs[0]
        p.style = doc.styles["Celda"]
        run = p.add_run(text.upper())
        run.font.size = Pt(7.4)
        run.bold = True
        run.font.color.rgb = INK3
        _char_spacing(run, 0.52)

    for r, row in enumerate(rows):
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            if r % 2 == 1:
                _shade(cell._tc.get_or_add_tcPr(), ZEBRA)
            _cell_margins(cell, top=75, bottom=75)
            p = cell.paragraphs[0]
            p.style = doc.styles["Celda"]
            add_rich(p, text, size=8.5, font=SANS)

    # Reparte el ancho disponible respetando la proporcion pedida.
    if widths:
        total = float(sum(widths))
        available = text_width(doc)
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = int(available * (w / total))

    _no_split(table)

    if caption:
        p = doc.add_paragraph(style=doc.styles["Pie"])
        add_rich(p, caption, size=8.2, color=INK3, font=SANS)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def _marker(paragraph, text, color, size):
    run = paragraph.add_run(text + "\t")
    run.font.name = SANS
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = True
    return run


def add_bullet(doc, text, level=1):
    if level == 1:
        p = doc.add_paragraph(style=doc.styles["Vineta"])
        _marker(p, "•", ACC, 9.5)
        add_rich(p, text)
    else:
        p = doc.add_paragraph(style=doc.styles["Vineta2"])
        _marker(p, "–", INK3, 9.0)
        add_rich(p, text, size=9.4)
    return p


def add_numbered(doc, items):
    """Lista numerada que siempre reinicia en 1."""
    last = None
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style=doc.styles["ListaNum"])
        _marker(p, "%d." % i, ACC, 9.5)
        if isinstance(item, (list, tuple)):
            add_rich(p, item[0])
            for child in item[1]:
                last = add_bullet(doc, child, level=2)
        else:
            add_rich(p, item)
            last = p
    return last


def close_list(last):
    """Separa la lista del parrafo siguiente sin insertar uno vacio."""
    if last is not None:
        last.paragraph_format.space_after = Pt(9)


def render_blocks(doc, blocks):
    for block in blocks:
        if isinstance(block, str):
            add_rich(doc.add_paragraph(), block)
            continue

        kind = block[0]

        if kind == "sub":
            # Sin size/color: los pone el estilo SubEnfoque.
            add_rich(doc.add_paragraph(style=doc.styles["SubEnfoque"]), block[1])
        elif kind in ("h1", "h2", "h3"):
            doc.add_heading(block[1], level=int(kind[1]))
        elif kind == "pagebreak":
            doc.add_page_break()
        elif kind == "lead":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(11)
            add_rich(p, block[1], size=11.4, color=INK2)
        elif kind == "bullets":
            last = None
            for item in block[1]:
                if isinstance(item, (list, tuple)):
                    add_bullet(doc, item[0], level=1)
                    for child in item[1]:
                        last = add_bullet(doc, child, level=2)
                else:
                    last = add_bullet(doc, item, level=1)
            close_list(last)
        elif kind == "numbers":
            close_list(add_numbered(doc, block[1]))
        elif kind == "code":
            add_code_block(doc, block[1])
        elif kind == "table":
            add_table(doc, block[1], block[2],
                      widths=block[3] if len(block) > 3 else None,
                      caption=block[4] if len(block) > 4 else None)
        elif kind == "box":
            add_box(doc, block[1], block[2], block[3])
        elif kind == "caption":
            add_rich(doc.add_paragraph(style=doc.styles["Pie"]), block[1],
                     size=8.2, color=INK3, font=SANS)
        else:
            raise ValueError("Bloque desconocido: %r" % (kind,))


# ------------------------------------------------------------------ modulo ---

FICHA_FIELDS = (
    ("objetivo", "Objetivo de aprendizaje"),
    ("duracion", "Duración estimada"),
    ("dependencias", "Requiere haber cursado"),
)


def render_ficha(doc, module):
    rows = [(label, module[key]) for key, label in FICHA_FIELDS if module.get(key)]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _table_borders(table, {
        "top": (HAIR, SZ_HAIR), "bottom": (HAIR, SZ_HAIR),
        "insideH": (HAIR, SZ_HAIR),
        "left": None, "start": None, "right": None, "end": None, "insideV": None,
    })

    available = text_width(doc)
    for i, (label, value) in enumerate(rows):
        left, right = table.rows[i].cells
        left.width = int(available * 0.26)
        right.width = int(available * 0.74)
        _cell_margins(left, top=90, bottom=90, start=0, end=110)
        _cell_margins(right, top=90, bottom=90, start=0, end=0)

        lp = left.paragraphs[0]
        lp.style = doc.styles["Etiqueta"]
        run = lp.add_run(label.upper())
        _char_spacing(run, 0.52)

        rp = right.paragraphs[0]
        rp.style = doc.styles["Celda"]
        add_rich(rp, value, size=9.2, font=SERIF)

    _no_split(table)


def render_module(doc, module):
    doc.add_page_break()

    heading = doc.add_heading("Módulo %s. %s" % (module["num"], module["title"]), level=2)
    heading.paragraph_format.space_before = Pt(0)
    for run in heading.runs:
        _char_spacing(run, -0.22)

    if module.get("tagline"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_rich(p, module["tagline"], size=10.4, color=INK3, font=SANS)

    render_ficha(doc, module)

    def section(title):
        p = doc.add_paragraph(style=doc.styles["SeccionModulo"])
        run = p.add_run(title.upper())
        _char_spacing(run, 0.56)

    section("Contenidos")
    add_numbered(doc, module["contenidos"])

    section("Enfoque didáctico")
    render_blocks(doc, module["enfoque"])

    section("Ejercicio práctico")
    render_blocks(doc, module["ejercicio"])

    section("Criterios de evaluación")
    for item in module["evaluacion"]:
        add_bullet(doc, item)


def part_stats(part):
    nums = [m["num"] for m in part["modules"]]
    hours = 0.0
    for module in part["modules"]:
        found = re.search(r"[\d.]+", module["duracion"])
        if found:
            hours += float(found.group())
    pretty = ("%.1f" % hours).rstrip("0").rstrip(".")
    span = "MÓDULO %d" % nums[0] if len(nums) == 1 else "MÓDULOS %d – %d" % (nums[0], nums[-1])
    return "%s   ·   %s HORAS" % (span, pretty)


def add_rule(doc, hex_color, width_cm, thickness_pt=2.4, space_after=10):
    """Barra de color solida (equivalente de .cover .rule, 34mm x 3px)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _table_borders(table, {e: None for e in
                           ("top", "left", "start", "bottom", "right", "end",
                            "insideH", "insideV")})

    # Altura exacta: si no, la altura minima de fila la engorda varias veces.
    trPr = table.rows[0]._tr.get_or_add_trPr()
    height = OxmlElement("w:trHeight")
    height.set(qn("w:val"), str(int(round(thickness_pt * 20))))
    height.set(qn("w:hRule"), "exact")
    trPr.append(height)

    cell = table.cell(0, 0)
    cell.width = Cm(width_cm)
    _shade(cell._tc.get_or_add_tcPr(), hex_color)
    _cell_margins(cell, top=0, bottom=0, start=0, end=0)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(thickness_pt)
    run = p.add_run(" ")
    run.font.size = Pt(1)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(space_after)
    spacer.add_run().font.size = Pt(1)


def render_part(doc, part):
    doc.add_page_break()
    for _ in range(3):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(part_stats(part))
    run.font.name = MONO
    run.font.size = Pt(8.4)
    run.font.color.rgb = INK3

    label = part["kicker"].replace("PARTE", "Parte")
    heading = doc.add_heading("%s · %s" % (label, part["title"]), level=1)
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(8)
    for run in heading.runs:
        _char_spacing(run, -0.4)

    add_rule(doc, ACC_HEX, 3.4, space_after=9)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(4)
    add_rich(intro, part["intro"], size=10.6, color=INK2)


# --------------------------------------------------------------- documento ---

def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)     # A4
    section.page_height = Mm(297)
    section.top_margin = Mm(22)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(12)
    section.footer_distance = Mm(11)
    return section


def add_footer(section, label):
    # La portada va limpia, como en el kit de prensa.
    section.different_first_page_header_footer = True

    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(label + "   ·   ")
    run.font.name = MONO
    run.font.size = Pt(7.6)
    run.font.color.rgb = INK3
    page = _field(p, "PAGE")
    page.font.name = MONO
    page.font.size = Pt(7.6)
    page.font.color.rgb = INK3


def add_spacer(doc, mm):
    """Hueco vertical exacto, en milimetros."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Mm(mm)
    p.add_run().font.size = Pt(1)
    return p


def add_cover(doc, meta):
    # CSS: .cover{padding-top:38mm} y .cover .mark{margin-bottom:26mm}
    add_spacer(doc, 34)

    # Distintivo de marca: rectangulo del acento con texto blanco.
    badge = doc.add_table(rows=1, cols=1)
    badge.alignment = WD_TABLE_ALIGNMENT.LEFT
    _table_borders(badge, {e: None for e in
                           ("top", "left", "start", "bottom", "right", "end",
                            "insideH", "insideV")})
    cell = badge.cell(0, 0)
    cell.width = Cm(2.55)
    _shade(cell._tc.get_or_add_tcPr(), ACC_HEX)
    _cell_margins(cell, top=90, bottom=110, start=170, end=170)
    bp = cell.paragraphs[0]
    bp.paragraph_format.space_after = Pt(0)
    bp.paragraph_format.line_spacing = 1.0
    brun = bp.add_run(meta["mark"])
    brun.font.name = SANS
    brun.font.size = Pt(11)
    brun.bold = True
    brun.font.color.rgb = WHITE
    _char_spacing(brun, -0.22)

    add_spacer(doc, 26)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.14
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(meta["title"])
    run.font.name = SANS
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = INK
    _char_spacing(run, -0.65)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(meta["subtitle"])
    run.font.name = SANS
    run.font.size = Pt(14)
    run.font.color.rgb = INK3
    _char_spacing(run, -0.2)

    add_spacer(doc, 9)                      # CSS: .cover h1{margin-bottom:9mm}
    add_rule(doc, ACC_HEX, 3.4, space_after=0)
    add_spacer(doc, 8)                      # CSS: .cover .rule{margin-bottom:8mm}

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(meta["who"])
    run.font.name = SANS
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = INK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(meta["who_sub"])
    run.font.name = SANS
    run.font.size = Pt(9.5)
    run.font.color.rgb = INK3

    add_spacer(doc, 16)                     # CSS: .cover .meta{margin-top:16mm}

    first = doc.add_paragraph()
    first.paragraph_format.space_before = Pt(0)
    first.paragraph_format.space_after = Pt(3)
    _para_border(first._p.get_or_add_pPr(), "top", HAIR, SZ_HAIR, space=11)
    for i, line in enumerate(meta["cover_meta"]):
        p = first if i == 0 else doc.add_paragraph()
        # Las dos ultimas lineas son procedencia: van en su propio grupo.
        gap = 10 if i == len(meta["cover_meta"]) - 3 else 3
        p.paragraph_format.space_after = Pt(gap)
        p.paragraph_format.line_spacing = 1.35
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_rich(p, line, size=8.4, color=INK3, font=MONO)


def add_toc(doc):
    doc.add_page_break()
    doc.add_heading("Índice", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    add_rich(p, "Si abres este archivo en Word y los números de página no aparecen, "
                "selecciona el índice y pulsa F9 para actualizarlo.",
             size=8.6, color=INK3, italic=True, font=SANS)
    _field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u')


def build_document():
    from content import program
    from content.syllabus import annexes, part1, part2, part3, part4, part5

    doc = Document()
    build_styles(doc)
    enable_hyphenation(doc)
    section = setup_page(doc)
    add_footer(section, program.META["footer"])

    add_cover(doc, program.META)
    add_toc(doc)

    doc.add_page_break()
    render_blocks(doc, program.FRONTMATTER)

    for part_module in (part1, part2, part3, part4, part5):
        for part in part_module.PARTS:
            render_part(doc, part)
            for module in part["modules"]:
                render_module(doc, module)

    for annex in annexes.ANEXOS:
        doc.add_page_break()
        doc.add_heading(annex["title"], level=1)
        render_blocks(doc, annex["blocks"])

    doc.save(DOCX_PATH)
    return DOCX_PATH


# --------------------------------------------------------------------- pdf ---

PS_EXPORT = r"""
$ErrorActionPreference = 'Stop'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
try {
    $doc = $word.Documents.Open('{DOCX}', $false, $true)
    foreach ($toc in $doc.TablesOfContents) { $toc.Update() }
    $doc.Fields.Update() | Out-Null
    foreach ($sec in $doc.Sections) {
        foreach ($hf in $sec.Footers) { $hf.Range.Fields.Update() | Out-Null }
    }
    $doc.Repaginate()
    $doc.ExportAsFixedFormat('{PDF}', 17, $false, 0, 0, 0, 0, 0, $true, $true, 1, $true, $true, $false)
    Write-Output ("PAGES=" + $doc.ComputeStatistics(2))
    $doc.Close($false)
} finally {
    $word.Quit()
    [System.Runtime.InteropServices.Marshal]::ReleaseComObject($word) | Out-Null
}
"""


def export_pdf(docx_path, pdf_path):
    script = PS_EXPORT.replace("{DOCX}", docx_path).replace("{PDF}", pdf_path)
    result = subprocess.run(
        ["powershell", "-NoProfile", "-NonInteractive", "-Command", script],
        capture_output=True, text=True,
    )
    print(result.stdout.strip())
    if result.returncode != 0:
        print(result.stderr.strip(), file=sys.stderr)
        raise SystemExit("Word no pudo exportar el PDF.")
    return pdf_path
